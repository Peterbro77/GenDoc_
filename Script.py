from docx import Document
from docx.shared import Pt, Inches
import os

doc = Document()

count = int(input("How many questions do you want to enter? "))

for i in range(count):
    question = input(f"Enter question {i+1}: ")

    print(f"Enter answer {i+1} (type 'END' on a new line to finish):")
    answer_lines = []
    while True:
        line = input()
        if line.strip().upper() == 'END':
            break
        answer_lines.append(line)
    answer = '\n'.join(answer_lines)

    # Add Question
    p = doc.add_paragraph()
    run_q = p.add_run(f"Q{i+1}: {question}\n")
    run_q.bold = True
    run_q.font.size = Pt(14)

    # Add Answer
    run_a = p.add_run(f"A{i+1}:\n{answer}\n\n")
    run_a.font.size = Pt(12)

    # Ask for image
    add_image = input("Do you want to insert an image for the output? (yes/no): ").strip().lower()
    if add_image == 'yes':
        image_path = input("Enter the full path to the image file: ").strip()
        if os.path.exists(image_path):
            # Add 'Output:' heading
            p_output = doc.add_paragraph()
            run_out = p_output.add_run("Output:\n")
            run_out.bold = True
            run_out.font.size = Pt(12)

            # Insert image (default width: 5 inches)
            try:
                doc.add_picture(image_path, width=Inches(5))
                doc.add_paragraph()  # Add a blank line
            except Exception as e:
                print(f"Couldn't insert image: {e}")
        else:
            print("Image file not found. Skipping image.")

docn = input("Enter the name for the file you want to put this into:")
doc.save(docn)
print(f"Saved to {docn}")
